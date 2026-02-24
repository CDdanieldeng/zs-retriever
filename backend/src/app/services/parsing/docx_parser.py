"""DOCX parser using python-docx."""

from io import BytesIO
from typing import BinaryIO

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.services.parsing.base import (
    Block,
    ImageBlock,
    Loc,
    TableBlock,
    TextBlock,
)


class DocxParser:
    """Parse DOCX files. Parent = section under heading."""

    def parse(self, content: bytes | BinaryIO, filename: str = "") -> list[Block]:
        """Parse DOCX content into blocks."""
        if isinstance(content, bytes):
            stream = BytesIO(content)
        else:
            stream = content
        doc = Document(stream)
        blocks: list[Block] = []
        heading_path: list[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                para = Paragraph(element, doc)
                style = para.style.name if para.style else ""
                if style.startswith("Heading"):
                    level = 1
                    for i, c in enumerate(style):
                        if c.isdigit():
                            level = int(style[i])
                            break
                    heading_path = heading_path[: level - 1] + [para.text.strip() or f"Section {level}"]
                    if para.text.strip():
                        blocks.append(
                            TextBlock(
                                content=para.text.strip(),
                                loc=Loc(heading_path=heading_path.copy()),
                                metadata={"style": style},
                            )
                        )
                else:
                    if para.text.strip():
                        blocks.append(
                            TextBlock(
                                content=para.text.strip(),
                                loc=Loc(heading_path=heading_path.copy()),
                            )
                        )
            elif tag == "tbl":
                table = Table(element, doc)
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    blocks.append(
                        TableBlock(
                            content="\n".join(rows),
                            loc=Loc(heading_path=heading_path.copy()),
                        )
                    )

        # Extract inline images from document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    img_part = rel.target_part
                    img_bytes = img_part.blob
                    blocks.append(
                        ImageBlock(
                            content="[Image]",
                            loc=Loc(heading_path=heading_path.copy()),
                            image_bytes=img_bytes,
                        )
                    )
                except Exception:
                    pass

        # If no headings, use single section
        if not blocks and doc.paragraphs:
            for para in doc.paragraphs:
                if para.text.strip():
                    blocks.append(
                        TextBlock(
                            content=para.text.strip(),
                            loc=Loc(heading_path=[]),
                        )
                    )
        if not blocks:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    blocks.append(
                        TableBlock(
                            content="\n".join(rows),
                            loc=Loc(heading_path=[]),
                        )
                    )
        return blocks
